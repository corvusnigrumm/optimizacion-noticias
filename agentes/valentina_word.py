"""
valentina_word.py — Agente Valentina (versión Word)

Aplica negrillas inteligentes basadas en los patrones editoriales
extraídos de "Recolección de notas de Natalia" y exporta un .docx
con estilos Heading 1, Heading 2 y Bold runs formateados.
"""
import json
import re
import os
from groq import Groq
from huggingface_hub import InferenceClient
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


MOLDE_NEGRILLAS = """
Eres Valentina, experta editora SEO de medios colombianos. Conoces a la perfección el estilo editorial de "La Redacción".
Analizarás el texto de una noticia y devolverás un JSON con dos listas:
  1. "subtitulos": lista de strings que son los subtítulos/secciones del artículo (H2). Se reconocen como preguntas en el texto o frases cortas que encabezan una sección.
  2. "frases": lista de frases LITERALES del cuerpo de la noticia que deben ir en negrilla.

REGLAS DE SUBTÍTULOS (H2):
- Son preguntas directas como "¿Por qué...?", "¿Cómo...?", "¿Cuándo...?"
- O encabezados de sección como "Paso a paso para...", "Los N beneficios de...", "Valor nutricional: ..."
- Cópialos TAL CUAL aparecen en el texto, sin modificar ni una letra.

REGLAS CRÍTICAS DE NEGRILLAS (basadas en patrones editoriales reales):
1. APERTURA IMPACTANTE: La frase de mayor impacto del primer párrafo va en negrilla (no el título completo, sino la afirmación central).
2. FUENTES Y CITAS: Siempre el nombre de la institución/fuente (Harvard, OMS, MedlinePlus, Cenor, etc.) + la conclusión clave que avala. Ej: "Harvard Health Publishing" y "los aminoácidos contribuyen a mantener la piel más firme".
3. TÉRMINOS TÉCNICOS: Palabras especializadas la primera vez que aparecen: colágeno, magnetrón, D-limoneno, Feng Shui, etileno, etc.
4. ADVERTENCIAS Y RIESGOS: Consecuencias negativas concretas. Ej: "puede provocar arcos eléctricos", "los billetes se deterioran".
5. INSTRUCCIONES/PASOS: En listas de pasos, el encabezado o acción principal de cada punto. Ej: "Usar solo trozos pequeños, lisos y sin arrugas".
6. DATOS NUMÉRICOS CLAVE: Cifras, porcentajes, tiempos concretos. Ej: "tres o cuatro días", "2.5 cm de distancia", "entre 8 y 9 gramos de proteína".
7. DENSIDAD: Entre 10 y 16 frases en negrilla. Máx 2 por párrafo. Nunca en párrafos de transición.
8. LONGITUD: Entre 4 y 15 palabras por frase. Nunca una oración completa de más de 20 palabras.
9. PROHIBIDO: Frases genéricas sin valor ("El artículo habla de..."), el título principal, pies de foto, nombres de autores.

Responde SOLO con el JSON, sin explicaciones:
{"subtitulos": ["subtítulo 1", "subtítulo 2", ...], "frases": ["frase literal 1", "frase literal 2", ...]}
"""


class ValentinaWord:
    """
    Agente que aplica negrillas usando el molde editorial de Natalia
    y exporta el resultado como documento Word (.docx).
    """

    def __init__(self):
        self.client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        self.hf_client = InferenceClient(api_key=os.getenv("HF_TOKEN"))

    def _llamar_llm(self, texto_crudo):
        """Llama al LLM con el molde de Natalia para obtener subtítulos y frases."""
        print("\n[ValentinaWord] Analizando texto con molde editorial: ", end="", flush=True)
        try:
            completion = self.client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": MOLDE_NEGRILLAS},
                    {"role": "user", "content": f"Texto de la noticia:\n\n{texto_crudo}"}
                ],
                temperature=0.3,
                max_completion_tokens=2048,
                top_p=0.95,
                response_format={"type": "json_object"},
                stream=True,
                stop=None
            )
            content = ""
            for chunk in completion:
                chunk_text = chunk.choices[0].delta.content or ""
                print(chunk_text, end="", flush=True)
                content += chunk_text
            print()

            # Limpiar bloques <think> o explicaciones antes del JSON
            content = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL)
            match = re.search(r"\{[\s\S]*\}", content)
            if match:
                content = match.group(0)

            data = json.loads(content)
            subtitulos = data.get("subtitulos", [])
            frases = data.get("frases", [])
            print(f"[ValentinaWord] ✅ {len(subtitulos)} H2 y {len(frases)} negrillas detectadas.")
            return subtitulos, frases

        except json.JSONDecodeError as e:
            print(f"\n[ValentinaWord] ⚠️ JSON inválido: {e}. Reintentando con HF...")
            return self._llamar_hf(texto_crudo)
        except Exception as e:
            if "RateLimitError" in type(e).__name__ or "429" in str(e):
                print("\n[ValentinaWord] ⚠️ Rate limit Groq. Usando Hugging Face...")
                return self._llamar_hf(texto_crudo)
            else:
                print(f"\n[ValentinaWord] ❌ Error: {e}")
                return [], []

    def _llamar_hf(self, texto_crudo):
        """Fallback a Hugging Face si Groq falla."""
        try:
            response = self.hf_client.chat.completions.create(
                model="Qwen/Qwen2.5-72B-Instruct",
                messages=[
                    {"role": "system", "content": MOLDE_NEGRILLAS},
                    {"role": "user", "content": f"Texto de la noticia:\n\n{texto_crudo}"}
                ],
                response_format={"type": "json_object"},
                max_tokens=2000
            )
            data = json.loads(response.choices[0].message.content)
            subtitulos = data.get("subtitulos", [])
            frases = data.get("frases", [])
            print(f"[ValentinaWord] ✅ HF: {len(subtitulos)} H2 y {len(frases)} negrillas.")
            return subtitulos, frases
        except Exception as e:
            print(f"[ValentinaWord] ❌ Error HF: {e}")
            return [], []

    def _enriquecer_con_busquedas(self, texto_crudo, busquedas_activas):
        """
        Añade las búsquedas activas de Google como contexto adicional
        para que el LLM priorice términos que la gente realmente busca.
        """
        if not busquedas_activas:
            return texto_crudo
        contexto = (
            "\n\n[BÚSQUEDAS ACTIVAS EN GOOGLE (Colombia)] "
            "Estos son los términos que la gente busca activamente. "
            "Prioriza resaltar en negrilla los que aparezcan en el texto:\n"
            + ", ".join(busquedas_activas[:20])
        )
        return texto_crudo + contexto

    def _construir_docx(self, titulo, lineas, subtitulos_set, frases, tags_seo=None):
        """
        Construye el documento Word:
        - Título → Heading 1
        - Subtítulos detectados → Heading 2
        - Resto → Normal con runs en Bold donde corresponda
        - Tags SEO → tabla al final (si se proporcionan)
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # --- Estilo global: fuente Georgia 12pt ---
        style = doc.styles["Normal"]
        style.font.name = "Georgia"
        style.font.size = Pt(12)

        # Configurar estilos Heading 1 y Heading 2
        h1_style = doc.styles["Heading 1"]
        h1_style.font.name = "Georgia"
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        h2_style = doc.styles["Heading 2"]
        h2_style.font.name = "Georgia"
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

        # Preparar regex de negrillas (ordenar por longitud desc para evitar matches parciales)
        frases_sorted = sorted(frases, key=len, reverse=True)
        patrones = [(f, re.compile(re.escape(f), re.IGNORECASE)) for f in frases_sorted if f.strip()]

        def aplicar_bold_en_parrafo(parrafo_word, texto):
            """Divide el texto en runs, marcando en bold las frases detectadas."""
            segmentos = []
            restante = texto
            while restante:
                mejor_match = None
                mejor_inicio = len(restante)
                for frase, patron in patrones:
                    m = patron.search(restante)
                    if m and m.start() < mejor_inicio:
                        mejor_match = m
                        mejor_inicio = m.start()
                if mejor_match:
                    if mejor_inicio > 0:
                        segmentos.append((restante[:mejor_inicio], False))
                    segmentos.append((mejor_match.group(), True))
                    restante = restante[mejor_match.end():]
                else:
                    segmentos.append((restante, False))
                    break

            for texto_seg, es_bold in segmentos:
                run = parrafo_word.add_run(texto_seg)
                run.bold = es_bold
                run.font.name = "Georgia"
                run.font.size = Pt(12)

        # --- Construir el documento línea por línea ---
        primera_linea = True
        for linea in lineas:
            linea = linea.strip()
            if not linea:
                continue

            linea_norm = linea.replace("\xa0", " ").strip()

            # Primera línea no vacía → H1
            if primera_linea:
                doc.add_heading(linea_norm, level=1)
                primera_linea = False
                continue

            # Subtítulo → H2
            es_h2 = any(
                linea_norm.lower() == s.replace("\xa0", " ").strip().lower()
                for s in subtitulos_set
            )
            if es_h2:
                doc.add_heading(linea_norm, level=2)
                continue

            # Párrafo normal con negrillas aplicadas
            parrafo = doc.add_paragraph(style="Normal")
            parrafo.paragraph_format.space_after = Pt(8)
            aplicar_bold_en_parrafo(parrafo, linea_norm)

        # --- Sección de Tags SEO al final ---
        if tags_seo:
            doc.add_page_break()
            doc.add_heading("Tags SEO", level=2)
            p_intro = doc.add_paragraph(style="Normal")
            p_intro.add_run("Búsquedas reales verificadas en Google Colombia (Google Suggest / Trends):")
            p_intro.paragraph_format.space_after = Pt(6)

            # Tabla: Tag | Tipo
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = "Table Grid"
            # Encabezados
            hdr = tabla.rows[0].cells
            for celda, texto_hdr in zip(hdr, ["Tag", "Tipo"]):
                celda.text = texto_hdr
                for run in celda.paragraphs[0].runs:
                    run.bold = True
                    run.font.name = "Georgia"
                    run.font.size = Pt(11)

            # Filas con los tags
            for tag_info in tags_seo:
                if isinstance(tag_info, dict):
                    tag_txt = tag_info.get("tag", "")
                    tipo_txt = tag_info.get("tipo", "Tendencia verificada")
                elif isinstance(tag_info, str):
                    tag_txt = tag_info
                    tipo_txt = "Tendencia verificada"
                else:
                    continue
                fila = tabla.add_row().cells
                fila[0].text = tag_txt
                fila[1].text = tipo_txt
                for celda in fila:
                    for run in celda.paragraphs[0].runs:
                        run.font.name = "Georgia"
                        run.font.size = Pt(11)

            print(f"[ValentinaWord] 🏷️  {len(tags_seo)} tags SEO añadidos al Word.")

        return doc

    def generar_docx(self, texto_crudo, ruta_salida, busquedas_activas=None, tags_seo=None):
        """
        Pipeline completo:
        1. Enriquece el texto con búsquedas activas de Google (si se proveen)
        2. Llama al LLM Qwen con el molde de Natalia
        3. Construye y guarda el .docx con negrillas, H1/H2 y tabla de Tags SEO

        Args:
            texto_crudo: El artículo en texto plano
            ruta_salida: Ruta donde guardar el .docx resultante
            busquedas_activas: Lista de términos de Google Suggest (opcional)
            tags_seo: Lista de dicts {tag, tipo, justificacion} generados por Pipe (opcional)

        Returns:
            ruta_salida si fue exitoso, None si falló
        """
        print("[ValentinaWord] 📝 Iniciando generación de Word con negrillas editoriales (Qwen)...")

        # 1. Enriquecer con búsquedas activas
        texto_enriquecido = self._enriquecer_con_busquedas(texto_crudo, busquedas_activas or [])

        # 2. Obtener subtítulos y frases del LLM (Qwen)
        subtitulos, frases = self._llamar_llm(texto_enriquecido)

        # 3. Procesar líneas del texto original (sin el contexto añadido)
        lineas = texto_crudo.split("\n")

        # 4. Construir el documento con tags SEO al final
        doc = self._construir_docx(
            titulo=lineas[0].strip() if lineas else "Artículo",
            lineas=lineas,
            subtitulos_set=subtitulos,
            frases=frases,
            tags_seo=tags_seo
        )

        # 5. Guardar
        os.makedirs(os.path.dirname(os.path.abspath(ruta_salida)), exist_ok=True)
        doc.save(ruta_salida)
        print(f"[ValentinaWord] ✅ Documento Word guardado en: {ruta_salida}")
        return ruta_salida
