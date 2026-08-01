import docx

doc = docx.Document(r'C:\Users\photo\Downloads\DESCARGAS\DESCARGAS\PROYECTOS\Optimización Noticias\Recolección de notas de Natalia.docx')
print(f"Total párrafos: {len(doc.paragraphs)}")

notas = []
nota_actual = []
titulo_actual = ""

for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if not txt:
        continue
    # Detectar separador tipo "Nota 1", "Nota 2", etc.
    stripped = txt.replace("Nota ", "").strip()
    if txt.startswith("Nota ") and stripped.isdigit():
        if nota_actual:
            notas.append({"titulo": titulo_actual, "parrafos": nota_actual})
        nota_actual = []
        titulo_actual = txt
    else:
        bold_parts = [r.text for r in para.runs if r.bold and r.text.strip()]
        nota_actual.append({"style": para.style.name, "text": txt, "bolds": bold_parts})

if nota_actual:
    notas.append({"titulo": titulo_actual, "parrafos": nota_actual})

print(f"Total notas encontradas: {len(notas)}")

for n in notas:
    parrafos = n["parrafos"]
    titulo_nota = parrafos[0]["text"] if parrafos else "(sin título)"
    print(f"\n{'='*60}")
    print(f"  {n['titulo']} => TÍTULO ARTÍCULO: {titulo_nota}")
    print(f"{'='*60}")

    # H2s: párrafos donde el texto completo está en negrilla
    h2s = []
    inline_bolds = []
    for p in parrafos[1:]:  # Saltar el título (primer párrafo)
        if not p["bolds"]:
            continue
        texto_bold_junto = "".join(p["bolds"]).strip().replace("\xa0", " ")
        texto_parrafo = p["text"].strip().replace("\xa0", " ")
        if texto_bold_junto == texto_parrafo or len(texto_bold_junto) >= len(texto_parrafo) * 0.9:
            h2s.append(p)
        else:
            inline_bolds.append(p)

    print(f"  H2/Subtítulos en negrilla completa ({len(h2s)}):")
    for h in h2s:
        print(f"    - {h['text']}")

    print(f"  Negrillas INLINE dentro de párrafo ({len(inline_bolds)}):")
    for b in inline_bolds[:8]:
        print(f"    PÁRRAFO: {b['text'][:100]}")
        print(f"    >> BOLD: {b['bolds']}")

# ---- ANÁLISIS GLOBAL DE PATRONES ----
print("\n\n" + "="*60)
print("ANÁLISIS DE PATRONES GLOBALES DE NEGRILLAS")
print("="*60)

h2_ejemplos = []
inline_ejemplos = []

for n in notas:
    for p in n["parrafos"]:
        if not p["bolds"]:
            continue
        texto_bold = "".join(p["bolds"]).strip().replace("\xa0", " ")
        texto = p["text"].strip().replace("\xa0", " ")
        if len(texto_bold) >= len(texto) * 0.85:
            h2_ejemplos.append(texto)
        else:
            inline_ejemplos.append({
                "parrafo": texto[:120],
                "bolds": p["bolds"]
            })

print(f"\nTotal H2/Subtítulos en negrita completa: {len(h2_ejemplos)}")
print(f"Total negrillas inline: {len(inline_ejemplos)}")

print("\n--- MUESTRA DE H2s (subtítulos en negrita) ---")
for e in h2_ejemplos[:20]:
    print(f"  > {e}")

print("\n--- MUESTRA DE NEGRILLAS INLINE (frases dentro de párrafos) ---")
for e in inline_ejemplos[:20]:
    print(f"  PÁRRAFO: {e['parrafo']}")
    print(f"  BOLD: {e['bolds']}")
    print()
